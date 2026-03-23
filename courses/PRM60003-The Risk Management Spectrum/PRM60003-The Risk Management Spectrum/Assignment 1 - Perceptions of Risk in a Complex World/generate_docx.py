from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Title
title = doc.add_heading('Perceptions of Risk in a Complex World: Evaluating the Impact of Risk Perception on Project Implementation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Abstract
abstract_heading = doc.add_heading('Abstract', 1)
abstract = doc.add_paragraph(
    'This essay critically examines how people\'s perceptions of risk influence the successful implementation of projects in complex operating environments. '
    'Drawing on established risk perception theories and analyzing notable project failures, this paper argues that cognitive biases, organizational culture, '
    'and communication breakdowns in risk perception significantly undermine project outcomes. Through case studies of the NASA Challenger disaster and the '
    'Boeing 737 MAX crisis, this essay demonstrates that effective risk management requires not only technical assessment but also deep understanding of '
    'human psychological and social factors in risk interpretation.'
)

# Keywords
keywords = doc.add_paragraph()
keywords.add_run('Keywords: ').bold = True
keywords.add_run('risk perception, complex projects, cognitive bias, organizational failure, project management')

doc.add_page_break()

# Section 1: Introduction
doc.add_heading('1. Introduction', 1)
p1 = doc.add_paragraph(
    'Risk is an inherent element of all projects, particularly those operating within complex environments characterized by uncertainty, interdependence, '
    'and dynamic stakeholder relationships (Bassi, 2024). While traditional risk management frameworks emphasize quantitative assessment and probabilistic '
    'analysis, a critical yet often overlooked dimension is how individuals and organizations '
)
p1.add_run('perceive').italic = True
p1.add_run(
    ' risk. Risk perception—the subjective judgment people make about the characteristics and severity of a risk—profoundly influences decision-making '
    'processes and, consequently, project outcomes (Slovic, 1987).\n\n'
    'This essay evaluates the impact of risk perception on project implementation in complex operating environments. It argues that discrepancies between '
    'objective risk assessment and subjective risk perception create vulnerabilities that can lead to catastrophic project failures. Through analysis of '
    'theoretical frameworks and empirical case studies, this paper demonstrates that successful project implementation requires integrating psychological '
    'insights into risk management practices.'
)

# Section 2
doc.add_heading('2. Theoretical Framework: Understanding Risk Perception', 1)

doc.add_heading('2.1 Defining Risk Perception', 2)
doc.add_paragraph(
    'Risk perception encompasses the judgments people make about the likelihood and severity of negative outcomes associated with particular hazards '
    '(Slovic, 1987). Unlike objective risk assessment, which relies on statistical probability and expected value calculations, risk perception is influenced '
    'by psychological, social, and cultural factors. Slovic\'s seminal work identified that laypeople\'s risk perceptions often diverge systematically from '
    'expert assessments, driven by factors such as dread, familiarity, and perceived control.'
)

doc.add_heading('2.2 Cognitive Biases in Risk Perception', 2)
doc.add_paragraph(
    'Human cognition is subject to systematic biases that distort risk perception. Kahneman and Tversky\'s (1979) prospect theory demonstrates that individuals '
    'evaluate potential losses and gains asymmetrically, exhibiting loss aversion and overweighting small probabilities. In project contexts, these biases manifest as:'
)

biases = [
    ('Optimism bias', 'The tendency to believe that projects will proceed more smoothly than statistically probable (Flyvbjerg, 2006).'),
    ('Confirmation bias', 'Selective attention to information supporting pre-existing risk assessments while dismissing contradictory evidence.'),
    ('Groupthink', 'The desire for conformity within teams leading to suppression of dissenting risk opinions (Janis, 1982).'),
    ('Availability heuristic', 'The tendency to judge risk probability based on how easily similar incidents come to mind.')
]

for bias, desc in biases:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{bias}: ').bold = True
    p.add_run(desc)

doc.add_heading('2.3 Organizational Dimensions of Risk Perception', 2)
doc.add_paragraph(
    'Organizations develop collective risk perceptions through shared narratives, routines, and power structures (Bassi, 2024). In hierarchical organizations, '
    'risk information may be filtered or distorted as it travels upward, creating what Vaughan (1996) termed "structural secrecy"—a normalization of deviance '
    'where risky practices become accepted through gradual accommodation.\n\n'
    'Organizational culture profoundly shapes what risks are visible and how they are interpreted. The "production imperative"—the pressure to meet deadlines '
    'and budgets—can systematically distort risk perception by framing safety concerns as obstacles to efficiency rather than legitimate project constraints.'
)

# Section 3
doc.add_heading('3. Case Study Analysis', 1)

doc.add_heading('3.1 NASA Challenger Disaster (1986)', 2)
doc.add_paragraph(
    'The Space Shuttle Challenger explosion represents a paradigmatic case of risk perception failure in complex projects. Despite warnings from engineers '
    'about O-ring seal vulnerabilities in cold weather, NASA management proceeded with the launch (Presidential Commission, 1986).'
)

p = doc.add_paragraph()
p.add_run('Risk Perception Factors:').bold = True

factors = [
    'Normalization of deviance: Previous successful launches with O-ring erosion created a false sense of security.',
    'Production pressure: The schedule-driven culture prioritized launch timelines over safety concerns.',
    'Communication breakdown: Technical risk information failed to translate into managerial risk perception.'
]

for factor in factors:
    doc.add_paragraph(factor, style='List Bullet')

doc.add_paragraph(
    'The Challenger case illustrates how organizational culture and power dynamics reshape risk perception. The disaster was not simply a technical failure '
    'but a failure of organizational risk perception—management had come to view O-ring erosion as a maintenance issue rather than a flight safety issue. '
    'This perceptual shift occurred gradually, with each successful launch reinforcing the belief that the system was safe despite accumulating evidence to the contrary.'
)

doc.add_heading('3.2 Boeing 737 MAX Crisis (2018-2019)', 2)
doc.add_paragraph(
    'The Boeing 737 MAX crashes demonstrate contemporary risk perception failures in commercial aviation projects. The Maneuvering Characteristics Augmentation '
    'System (MCAS) was designed to address aerodynamic instabilities, but its risk implications were inadequately perceived by both Boeing and regulators '
    '(House Committee, 2020).'
)

p = doc.add_paragraph()
p.add_run('Risk Perception Factors:').bold = True

factors2 = [
    'Cost-driven risk assessment: Competitive pressure led to compressed development timelines and underestimation of software risks.',
    'Regulatory capture: The FAA\'s delegation of certification authority created conflicts of interest.',
    'Information asymmetry: Pilots were not fully informed about MCAS functionality.'
]

for factor in factors2:
    doc.add_paragraph(factor, style='List Bullet')

doc.add_paragraph